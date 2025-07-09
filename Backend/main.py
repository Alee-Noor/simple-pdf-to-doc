import os
import tempfile
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from starlette.background import BackgroundTask
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import fitz  # PyMuPDF
from PIL import Image

# Set your Tesseract-OCR installation path (Windows-specific)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

app = FastAPI()

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

def cleanup_files(*paths):
    """Clean up temporary files"""
    for path in paths:
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except Exception as e:
                print(f"Error cleaning up {path}: {e}")

def pdf_to_docx_skip_errors(pdf_path: str, docx_path: str) -> None:
    """Convert PDF to DOCX with error handling"""
    doc = Document()
    pdf_document = None
    
    try:
        pdf_document = fitz.open(pdf_path)
        pages = convert_from_path(pdf_path, dpi=300)

        if len(pages) != len(pdf_document):
            raise RuntimeError("Page count mismatch between PDF and converted images")

        for page_number, (page_image, pdf_page) in enumerate(zip(pages, pdf_document), start=1):
            try:
                # Image extraction
                images = pdf_page.get_images(full=True)
                if images:
                    for img_index, img in enumerate(images):
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{image_ext}") as img_file:
                            img_file.write(image_bytes)
                            temp_image_path = img_file.name

                        if img_index == 0:
                            doc.add_picture(temp_image_path)
                            doc.add_paragraph()
                        os.remove(temp_image_path)

                # Text extraction
                text = pytesseract.image_to_string(page_image)
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                for para in paragraphs:
                    doc.add_paragraph(para)
                doc.add_page_break()

            except Exception as page_error:
                print(f"Skipping page {page_number}: {page_error}")
                continue

        doc.save(docx_path)
        if not os.path.exists(docx_path):
            raise RuntimeError("Failed to save DOCX file")

    except Exception as e:
        # Cleanup docx if conversion failed
        if os.path.exists(docx_path):
            os.remove(docx_path)
        raise
    finally:
        if pdf_document:
            pdf_document.close()

@app.post("/convert")
async def convert_pdf_to_docx(file: UploadFile = File(...)):
    """Handle PDF conversion endpoint"""
    temp_pdf_path = None
    temp_docx_path = None

    try:
        # Validate input
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(400, "Only PDF files are allowed")

        # Create temp PDF file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            content = await file.read()
            if not content:
                raise HTTPException(400, "Empty file uploaded")
            tmp_pdf.write(content)
            temp_pdf_path = tmp_pdf.name

        # Create temp DOCX path
        temp_docx_path = tempfile.mktemp(suffix=".docx")

        # Perform conversion
        pdf_to_docx_skip_errors(temp_pdf_path, temp_docx_path)

        # Verify output
        if not os.path.exists(temp_docx_path):
            raise HTTPException(500, "Conversion failed - no output generated")

        # Return response with cleanup task
        return FileResponse(
            temp_docx_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename="converted.docx",
            background=BackgroundTask(cleanup_files, temp_pdf_path, temp_docx_path)
        )
        
    except HTTPException:
        cleanup_files(temp_pdf_path, temp_docx_path)
        raise
    except Exception as e:
        cleanup_files(temp_pdf_path, temp_docx_path)
        raise HTTPException(500, f"Conversion failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    